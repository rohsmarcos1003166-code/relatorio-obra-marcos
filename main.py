import streamlit as st
import pandas as pd
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import io

# ==========================================
# CONFIGURAÇÕES DO SEU PROJETO (Marcos)
# ==========================================
# Link atualizado da sua nova planilha do Google
LINK_PLANILHA_CSV = "https://docs.google.com/spreadsheets/d/1xFpGJI7Bq3qj48hVSUO190qHE-N2kXg3r-BmQaSM1uk/export?format=csv"
TELEGRAM_TOKEN = "8568148429:AAGzu7zf-n-fGJnUpaNGVCLvQnsR2JxJ3fs"
TELEGRAM_CHAT_ID = "7668457919"

# ==========================================
# FUNÇÃO PRINCIPAL DO ROBÔ
# ==========================================
def gerar_e_enviar_relatorio():
    try:
        # 1. Lê a planilha atualizada do Google
        df = pd.read_csv(LINK_PLANILHA_CSV)
        
        # Limpa espaços invisíveis que possam existir nos nomes das colunas
        df.columns = df.columns.str.strip()

        # Mapeamento inteligente baseado exatamente nas colunas da sua nova planilha
        mapeamento = {
            'Material (A)': 'Material',
            'Unid (B)': 'Unid',
            'Mínimo (C)': 'Mínimo',
            'Estoque atual(saldo)': 'Atual',
            'Urgencia': 'Urgência',
            'Preço Ref (F)': 'Preço Ref',
            'Preço atual': 'Preço Atual',
            'Valor total': 'Valor Total'
        }
        df.rename(columns=mapeamento, inplace=True)

        # 2. Inicia o arquivo de Word
        doc = Document()
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tit = p_titulo.add_run("🏗️ Relatório Diário Prédio / Acompanhamento da Obra")
        run_tit.font.size = Pt(16)
        run_tit.font.bold = True
        
        doc.add_paragraph(f"Responsável Técnico: Marcos Balaur")
        doc.add_paragraph(f"Data de Emissão: {datetime.date.today().strftime('%d/%m/%Y')}")
        doc.add_paragraph("-" * 60)

        # 3. Análise de Alertas de Estoque
        doc.add_heading("🚨 Alertas de Estoque Crítico", level=2)
        alertas_gerados = 0
        
        for idx, linha in df.iterrows():
            # Puxa os dados tratando textos e números vazios
            atual = pd.to_numeric(linha.get('Atual', 0), errors='coerce')
            minimo = pd.to_numeric(linha.get('Mínimo', 0), errors='coerce')
            material = linha.get('Material', 'Item Desconhecido')
            urgencia = str(linha.get('Urgência', 'BAIXA')).strip().upper()
            
            if not pd.isna(atual) and not pd.isna(minimo) and atual <= minimo:
                alertas_gerados += 1
                if urgencia == "ALTA" or urgencia == "CRÍTICA":
                    p = doc.add_paragraph()
                    r = p.add_run(f"🔴 CRÍTICO: {material} está com saldo {atual:.0f} (Mínimo exigido: {minimo:.0f}). Urgência de Compra Máxima!")
                    r.font.bold = True
                else:
                    doc.add_paragraph(f"🟡 ATENÇÃO: {material} atingiu o limite mínimo ({atual:.0f}/{minimo:.0f}). Planejar reposição.")

        if alertas_gerados == 0:
            doc.add_paragraph("✅ Todos os níveis de materiais monitorados estão operando acima do estoque mínimo.")

        doc.add_paragraph("\n")

        # 4. Busca Automática de Fotos em qualquer parte da planilha
        doc.add_heading("📸 Comprovação Visual do Andamento", level=2)
        foto_anexada = False
        
        for col in df.columns:
            valores_coluna = df[col].dropna().astype(str).tolist()
            for valor in valores_coluna:
                valor_limpo = valor.strip()
                # Se encontrar qualquer célula começando com link de internet
                if valor_limpo.startswith("http://") or valor_limpo.startswith("https://"):
                    # Se for link do postimg, imgur, google drive ou contiver extensão de imagem
                    if any(ext in valor_limpo.lower() for ext in ['.jpg', '.jpeg', '.png', 'postimg', 'drive.google']):
                        try:
                            resposta_imagem = requests.get(valor_limpo, timeout=15)
                            if resposta_imagem.status_code == 200:
                                memoria_imagem = io.BytesIO(resposta_imagem.content)
                                doc.add_picture(memoria_imagem, width=Inches(4.5))
                                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                doc.add_paragraph("Legenda: Registro fotográfico anexado automaticamente via planilha.").alignment = WD_ALIGN_PARAGRAPH.CENTER
                                foto_anexada = True
                                break
                        except Exception:
                            continue
            if foto_anexada:
                break

        if not foto_anexada:
            doc.add_paragraph("[Nenhum link de foto ativo ou válido inserido na planilha para este período.]")

        doc.add_paragraph("\n")

        # 5. Tabela Geral de Materiais
        doc.add_heading("📋 Tabela Consolidada de Controle", level=2)
        
        # Filtra colunas de links para não poluírem a tabela visual do Word
        colunas_exibicao = [c for c in df.columns if 'http' not in str(c).lower() and '.jpg' not in str(c).lower()]
        
        tabela_dados = doc.add_table(rows=1, cols=len(colunas_exibicao))
        tabela_dados.style = 'Table Grid'
        
        hdr_cells = tabela_dados.rows[0].cells
        for i, nome_coluna in enumerate(colunas_exibicao):
            hdr_cells[i].text = str(nome_coluna)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
        for _, linha in df.iterrows():
            celulas_linha = tabela_dados.add_row().cells
            for i, col in enumerate(colunas_exibicao):
                valor_celula = linha[col]
                if ('Preço' in col or 'Valor' in col) and isinstance(valor_celula, (int, float)) and not pd.isna(valor_celula):
                    celulas_linha[i].text = f"R$ {valor_celula:,.2f}"
                else:
                    celulas_linha[i].text = str(valor_celula) if not pd.isna(valor_celula) else ""

        # 6. Salva o arquivo de Word e despacha para o Telegram
        nome_doc = "Relatorio_Automatico_Obra.docx"
        doc.save(nome_doc)

        url_telegram = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument"
        with open(nome_doc, "rb") as arquivo_word:
            payload_arquivos = {"document": arquivo_word}
            payload_dados = {
                "chat_id": TELEGRAM_CHAT_ID,
                "caption": "📊 *Relatório Atualizado da Obra!*\nMarcos, o relatório foi gerado com sucesso puxando os dados da sua planilha nova!"
            }
            resposta = requests.post(url_telegram, data=payload_dados, files=payload_arquivos)
            
        return resposta.status_code == 200

    except Exception as e:
        st.error(f"Erro na execução geral: {e}")
        return False

# ==========================================
# PAINEL DE INTERFACE DO STREAMLIT
# ==========================================
st.title("🚀 Central de Automação de Engenharia")
st.write("Conectado à Planilha do Google e integrado ao Telegram.")

st.markdown("---")

if st.button("Disparar Primeiro Teste Agora"):
    with st.spinner("Buscando dados e fotos na planilha..."):
        sucesso = gerar_e_enviar_relatorio()
    if sucesso:
        st.success("Perfeito Marcos! Relatório enviado com sucesso para o Telegram.")
    else:
        st.error("Erro ao gerar relatório. Verifique se a planilha está pública e acessível.")
